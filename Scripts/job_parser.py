import pandas as pd
import re
import os
import logging
import csv
from datetime import datetime
from collections import defaultdict
import pdfplumber
from docx import Document
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from settings import JD_FOLDER, PARSED_JD_FILE, KEYWORDS_FILE, LOG_FILE

from utils import setup_logger, log_exceptions

setup_logger(LOG_FILE)
logger = logging.getLogger(__name__)

@log_exceptions
def load_keywords():
    df = pd.read_csv(KEYWORDS_FILE)
    if 'Skills' not in df.columns or 'Keywords' not in df.columns:
        raise KeyError("Both 'Skills' and 'Keywords' columns are required.")

    skill_keywords = df['Skills'].dropna().str.lower().str.strip().str.split(',').explode().str.strip().tolist()
    requirement_keywords = df['Keywords'].dropna().str.lower().str.strip().str.split(',').explode().str.strip().tolist()

    return set(skill_keywords), set(requirement_keywords)

def extract_keywords(description, skill_keywords, requirement_keywords):
    description_lower = description.lower()
    skill_freq = defaultdict(int)
    requirement_freq = defaultdict(int)

    for skill in skill_keywords:
        pattern = r'\b' + re.escape(skill) + r'\b'
        matches = re.findall(pattern, description_lower)
        if matches:
            skill_freq[skill] = len(matches)

    for req in requirement_keywords:
        if req in skill_freq:
            continue
        pattern = r'\b' + re.escape(req) + r'\b'
        matches = re.findall(pattern, description_lower)
        if matches:
            requirement_freq[req] = len(matches)

    return list(skill_freq.keys()), list(requirement_freq.keys()), skill_freq, requirement_freq

def extract_job_title(description):
    for line in description.splitlines():
        match = re.search(r'(?:title[:\-]?)\s*([^\n\r]{3,150})', line, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return "Unknown"

def extract_title_from_formatting(filepath):
    try:
        if filepath.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    words = page.extract_words(use_text_flow=True)
                    for word in words:
                        if float(word.get("size", 0)) > 10 and word.get("fontname", "").lower().startswith("bold"):
                            return word['text'].strip()
        elif filepath.endswith(".docx"):
            doc = Document(filepath)
            for para in doc.paragraphs:
                if para.style.name.lower().startswith("heading") or para.text.strip():
                    return para.text.strip()
        return "Unknown"
    except Exception as e:
        logger.error(f"Formatting-based title extraction failed: {e}")
        return "Unknown"

@log_exceptions
def get_next_job_id(output_file):
    if not os.path.exists(output_file):
        return "JOB101"

    df = pd.read_csv(output_file)
    if 'job_id' not in df.columns or df.empty:
        return "JOB101"
    existing_ids = df['job_id'].dropna().tolist()
    nums = [
        int(re.search(r'JOB(\d+)', jid).group(1))
        for jid in existing_ids if re.match(r'JOB\d+', jid)
    ]
    next_num = max(nums) + 1 if nums else 101
    return f"JOB{next_num}"

@log_exceptions
def save_to_csv(output_file, job_id, job_title, skill_freq, requirement_freq):
    file_exists = os.path.isfile(output_file)
    with open(output_file, mode='a', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        if not file_exists:
            writer.writerow(['job_id', 'job_title', 'skill', 'weight', 'timestamp'])

        timestamp = datetime.now().strftime("%Y-%m-%d")

        for skill, freq in skill_freq.items():
            writer.writerow([job_id, job_title, skill, freq, timestamp])

        for req, freq in requirement_freq.items():
            writer.writerow([job_id, job_title, req, freq, timestamp])

def is_duplicate_job(output_file, job_title):
    if not os.path.exists(output_file):
        return False
    df = pd.read_csv(output_file)
    return job_title.strip().lower() in df['job_title'].str.lower().str.strip().values

@log_exceptions
def process_job_description(description, filepath, output_file=PARSED_JD_FILE):
    skill_keywords, requirement_keywords = load_keywords()
    job_title = extract_job_title(description)

    if job_title == "Unknown" or job_title.lower() == "job title":
        job_title = extract_title_from_formatting(filepath)

    # ✅ Check for duplicate
    if is_duplicate_job(output_file, job_title):
        print(f"Skipped duplicate job: {job_title}")
        return

    job_id = get_next_job_id(output_file)
    skills, requirements, skill_freq, requirement_freq = extract_keywords(
        description, skill_keywords, requirement_keywords
    )

    save_to_csv(output_file, job_id, job_title, skill_freq, requirement_freq)
    print(f"Processed job '{job_title}' (ID: {job_id})")
@log_exceptions
def parse_jobs():
    for filename in os.listdir(JD_FOLDER):
        filepath = os.path.join(JD_FOLDER, filename)
        content = ""

        try:
            if filename.endswith(".txt"):
                with open(filepath, 'r', encoding='utf-8') as file:
                    content = file.read()

            elif filename.endswith(".docx"):
                doc = Document(filepath)
                content = "\n".join([para.text for para in doc.paragraphs])

            elif filename.endswith(".pdf"):
                with pdfplumber.open(filepath) as pdf:
                    content = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

            if content:
                process_job_description(content, filepath)
                print(f"Processed: {filename}")

        except Exception as e:
            logger.error(f"Failed to process {filename}: {e}")
            print(f"Failed to process {filename}: {e}")

if __name__ == "__main__":
    parse_jobs()



